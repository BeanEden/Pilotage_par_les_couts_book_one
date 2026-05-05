"""Generateur Word — Donnees essentielles Book One pour NotebookLM."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def h1(txt): doc.add_heading(txt, level=1)
def h2(txt): doc.add_heading(txt, level=2)
def h3(txt): doc.add_heading(txt, level=3)
def p(txt): doc.add_paragraph(txt)
def b(txt): r = doc.add_paragraph().add_run(txt); r.bold = True

# ============================================================
h1("BOOK ONE — Donnees essentielles du projet")
# ============================================================
p("Document de reference pour le pilotage par les couts du projet Book One MVP Mobile.")
p("Genere automatiquement a partir du code source de l'application Streamlit.")

# ============================================================
h1("1. FICHE PROJET")
# ============================================================
t = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
data = [
    ("Nom du projet", "Book One — Tous les livres a 1 EUR"),
    ("Type", "MVP Application Mobile (Flutter)"),
    ("Duree", "8 semaines (6 janvier — 28 fevrier 2025)"),
    ("Budget valide", "130 000 EUR"),
    ("CA previsionnel annuel", "180 000 EUR"),
    ("Equipe", "5 profils actifs (PM, BE, MOB, QA, FRL)"),
    ("Methodologie", "Agile — jalons S4 et S8"),
    ("Stack technique", "Flutter (mobile), Node.js (API), PostgreSQL, Stripe, Firebase"),
]
for i, (k, v) in enumerate(data):
    t.rows[i].cells[0].text = k
    t.rows[i].cells[1].text = v

# ============================================================
h1("2. EQUIPE & RESSOURCES HUMAINES")
# ============================================================

h2("2.1 Composition de l'equipe (9 profils initiaux)")
t = doc.add_table(rows=10, cols=6, style='Light Grid Accent 1')
headers = ["ID", "Poste", "Type contrat", "Statut", "Salaire brut annuel", "TJM employeur"]
for i, h in enumerate(headers): t.rows[0].cells[i].text = h
equipe = [
    ("PM", "Chef de projet", "CDI", "Cadre", "58 000 EUR", "~386 EUR/j"),
    ("TL", "Tech Lead", "CDI", "Cadre", "72 000 EUR", "~479 EUR/j"),
    ("BE", "Dev Back-End", "CDI", "ETAM", "52 000 EUR", "~338 EUR/j"),
    ("MOB", "Dev Mobile Flutter", "CDI", "ETAM", "55 000 EUR", "~358 EUR/j"),
    ("UX", "UX Designer", "CDI", "ETAM", "48 000 EUR", "~313 EUR/j"),
    ("QA", "QA Engineer", "CDI", "ETAM", "45 000 EUR", "~293 EUR/j"),
    ("STG", "Stagiaire Dev", "Stage", "Stagiaire", "7 800 EUR", "~52 EUR/j"),
    ("ALT", "Alternant Dev", "Alternance", "Alternant", "16 000 EUR", "~81 EUR/j"),
    ("FRL", "Freelance Paiement/Secu", "Freelance", "Freelance", "130 000 EUR", "~650 EUR/j"),
]
for i, row in enumerate(equipe):
    for j, val in enumerate(row):
        t.rows[i+1].cells[j].text = val

h2("2.2 Calcul du TJM")
p("Formule : TJM = (Salaire brut annuel / Jours ouvres) x (1 + Taux charges patronales)")
t = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(["Statut", "Charges patronales", "Jours ouvres/an"]):
    t.rows[0].cells[i].text = h
for i, (s, c, j) in enumerate([("Cadre", "45%", "218"), ("ETAM", "42%", "218"), ("Stagiaire/Alternant", "0-10%", "151-218"), ("Freelance", "0%", "200")]):
    t.rows[i+1].cells[0].text = s
    t.rows[i+1].cells[1].text = c
    t.rows[i+1].cells[2].text = j

h2("2.3 Levier d'optimisation RH")
p("Remplacement du Dev Back-End Senior (52 000 EUR) par un profil Junior (38 000 EUR).")
p("Economie estimee : 6 500 EUR sur la duree du projet (8 semaines).")
p("Impact : reduction du TJM de 338 EUR/j a 247 EUR/j — soit -27%.")

# ============================================================
h1("3. STRUCTURE BUDGETAIRE")
# ============================================================

h2("3.1 Enveloppes budgetaires")
t = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
for i, h in enumerate(["Poste", "Montant", "% du total"]):
    t.rows[0].cells[i].text = h
for i, (p2, m, pc) in enumerate([
    ("Ressources Humaines", "112 000 EUR", "86.2%"),
    ("Couts hors-RH", "7 200 EUR", "5.5%"),
    ("Provision risques", "10 800 EUR", "8.3%"),
    ("TOTAL BUDGET VALIDE", "130 000 EUR", "100%"),
]):
    t.rows[i+1].cells[0].text = p2
    t.rows[i+1].cells[1].text = m
    t.rows[i+1].cells[2].text = pc

h2("3.2 Detail des couts hors-RH (7 200 EUR)")
t = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
t.rows[0].cells[0].text = "Categorie"; t.rows[0].cells[1].text = "Montant"
cats = [("Infrastructure & Cloud", "1 800 EUR"), ("Licences & Logiciels", "700 EUR"),
        ("Donnees & APIs", "400 EUR"), ("Qualite & Tests", "800 EUR"),
        ("Legal & Conformite", "1 100 EUR"), ("Operationnel", "900 EUR"),
        ("Provision & Risques", "1 500 EUR")]
for i, (c, m) in enumerate(cats):
    t.rows[i+1].cells[0].text = c; t.rows[i+1].cells[1].text = m

h2("3.3 Ratios financiers cles")
p("- Ratio RH / Budget total : 86.2%")
p("- Ratio Hors-RH / Budget total : 5.5%")
p("- Reserve risques / Budget total : 8.3%")
p("- Cout reel S8 : 33 422 EUR vs Planifie : 31 590 EUR → Ecart : +5.8%")
p("- Budget consomme a S8 : 33 422 EUR / 130 000 EUR = 25.7% du budget valide")
p("- Facteur de derive moyen RH : x1.08 (ecart planifie vs reel)")

# ============================================================
h1("4. PLANNING & GANTT")
# ============================================================

h2("4.1 Phasage du projet (8 semaines)")
t = doc.add_table(rows=7, cols=4, style='Light Grid Accent 1')
for i, h in enumerate(["Phase", "Semaines", "Budget estime", "Livrables"]):
    t.rows[0].cells[i].text = h
phases = [
    ("P1 Cadrage & Architecture", "S1-S2", "5 200 EUR", "Backlog, stack, archi cloud"),
    ("P2 UX/UI & Architecture", "S1-S3", "12 800 EUR", "Maquettes HD, CI/CD, modele donnees"),
    ("P3 Dev MVP", "S3-S6", "32 400 EUR", "API + App mobile + Messagerie"),
    ("P4 Fonctionnalites Bonus", "S5-S7", "14 500 EUR", "Stripe, Geoloc, Invite, Remise"),
    ("P5 Tests & Recette", "S6-S8", "9 800 EUR", "QA fonctionnel + securite + UAT"),
    ("P6 Production", "S7-S8", "5 700 EUR", "Deploiement + Stores + Monitoring"),
]
for i, (ph, sem, bud, liv) in enumerate(phases):
    t.rows[i+1].cells[0].text = ph; t.rows[i+1].cells[1].text = sem
    t.rows[i+1].cells[2].text = bud; t.rows[i+1].cells[3].text = liv

h2("4.2 Liste des 28 taches (Gantt)")
t = doc.add_table(rows=29, cols=7, style='Light Grid Accent 1')
for i, h in enumerate(["ID", "Categorie", "Nom", "Ressource", "Semaine", "Duree", "Critique"]):
    t.rows[0].cells[i].text = h
taches = [
    (1,"Architecture","Definition architecture technique","TL",1,2,"Oui"),
    (2,"Architecture","Setup repo, CI/CD & environnements","TL",1,2,"Non"),
    (3,"Architecture","Specifications fonctionnelles MVP","STG",1,3,"Non"),
    (4,"Authentification","Maquettes UX auth + mode invite","UX",3,2,"Oui"),
    (5,"Authentification","API Auth (register, login, JWT)","BE",3,3,"Oui"),
    (6,"Authentification","Mode invite — acces lecture catalogue","MOB",5,2,"Non"),
    (7,"Catalogue","Modele de donnees livres + geoloc","BE",3,2,"Oui"),
    (8,"Catalogue","API catalogue — recherche & filtres geo","BE",5,3,"Oui"),
    (9,"Catalogue","Ecran carte — livres autour de moi","MOB",6,3,"Oui"),
    (10,"Catalogue","Ecran fiche livre detail","MOB",8,2,"Oui"),
    (11,"Catalogue","Favoris & Wishes (livres souhaites)","ALT",8,3,"Non"),
    (12,"Depot","Integration API ISBN (Open Library)","BE",5,2,"Oui"),
    (13,"Depot","Ecran scan ISBN (camera mobile)","MOB",6,3,"Oui"),
    (14,"Depot","Logique credit : 3 scans → 1 achat","BE",7,2,"Oui"),
    (15,"Paiement","Integration Stripe — paiement 1 EUR","FRL",9,3,"Oui"),
    (16,"Paiement","Ecran achat + recap commande","MOB",10,2,"Oui"),
    (17,"Paiement","Debit conditionnel (apres confirmation)","BE",11,2,"Oui"),
    (18,"Paiement","Audit securite flux paiement","FRL",12,2,"Oui"),
    (19,"Livraison","Messagerie in-app — convenir du RDV","BE",9,3,"Non"),
    (20,"Livraison","Ecran messagerie & confirmation RDV","MOB",11,2,"Non"),
    (21,"Livraison","Confirmation remise en main propre","MOB",12,2,"Oui"),
    (22,"Notifications","Service push notifications (FCM/APNs)","BE",7,3,"Non"),
    (23,"Notifications","Notifs : livre vendu + wish disponible","ALT",9,3,"Non"),
    (24,"QA","Tests unitaires back-end","ALT",8,4,"Non"),
    (25,"QA","Tests integration & E2E","QA",11,3,"Oui"),
    (26,"DevOps","Config staging / prod + monitoring","TL",5,2,"Non"),
    (27,"DevOps","Documentation technique","STG",9,4,"Non"),
    (28,"Lancement","Beta test interne + correction bugs","QA",13,2,"Oui"),
]
for i, (tid,cat,nom,res,sem,dur,crit) in enumerate(taches):
    t.rows[i+1].cells[0].text = str(tid); t.rows[i+1].cells[1].text = cat
    t.rows[i+1].cells[2].text = nom; t.rows[i+1].cells[3].text = res
    t.rows[i+1].cells[4].text = f"S{sem}"; t.rows[i+1].cells[5].text = f"{dur} sem"
    t.rows[i+1].cells[6].text = crit

# ============================================================
h1("5. KPIs & INDICATEURS DE PILOTAGE")
# ============================================================

h2("5.1 KPIs au jalon S4 (mi-parcours)")
t = doc.add_table(rows=7, cols=2, style='Light Grid Accent 1')
for i, (k,v) in enumerate([("Jalon","S4 — Semaine 4"),("Taches terminees","8 / 10"),
    ("Taches en retard","2"),("Taux achevement","33%"),
    ("Cout planifie","~15 800 EUR (cumule S1-S4)"),
    ("Cout reel","~16 200 EUR (cumule S1-S4)"),("Couts non planifies","650 EUR")]):
    t.rows[i].cells[0].text = k; t.rows[i].cells[1].text = v

h2("5.2 KPIs au jalon S8 (fin de projet)")
t = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
for i, (k,v) in enumerate([("Jalon","S8 — Semaine 8 (final)"),("Taches terminees","24 / 24"),
    ("Taches en retard","0"),("Taux achevement","100%"),
    ("Cout planifie","31 590 EUR"),("Cout reel","33 422 EUR"),
    ("Ecart budgetaire","+5.8% (derive maitrisee)"),("Couts non planifies","4 000 EUR")]):
    t.rows[i].cells[0].text = k; t.rows[i].cells[1].text = v

h2("5.3 Facteurs de derive hebdomadaire (reel vs planifie)")
p("Coefficients appliques au cout RH planifie pour obtenir le cout reel :")
p("S1: 0.95 | S2: 0.98 | S3: 1.00 | S4: 1.05 | S5: 1.15 | S6: 1.10 | S7: 1.08 | S8: 1.10")
p("Interpretation : les semaines S5-S8 montrent une acceleration des couts (+8 a 15%), liee a la complexite croissante des developpements (Stripe, tests E2E).")

h2("5.4 Meteo equipe (moral des collaborateurs)")
t = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
t.rows[0].cells[0].text = "Ressource"; t.rows[0].cells[1].text = "S4"; t.rows[0].cells[2].text = "S8"
meteo = [("PM — Chef de projet","3/3 Excellente","2/3 Correcte"),
         ("BE — Dev Back-End","2/3 Correcte","2/3 Correcte"),
         ("MOB — Dev Mobile","3/3 Excellente","1/3 Difficile"),
         ("QA — QA Engineer","2/3 Correcte","3/3 Excellente"),
         ("FRL — Freelance","3/3 Excellente","2/3 Correcte")]
for i, (r,s4,s8) in enumerate(meteo):
    t.rows[i+1].cells[0].text = r; t.rows[i+1].cells[1].text = s4; t.rows[i+1].cells[2].text = s8

# ============================================================
h1("6. GESTION DES RISQUES")
# ============================================================

h2("6.1 Risques survenus")
t = doc.add_table(rows=5, cols=6, style='Light Grid Accent 1')
for i, h in enumerate(["ID","Risque","Categorie","Impact","Statut","Cout"]):
    t.rows[0].cells[i].text = h
risques = [
    ("RS1","API ISBN : limite de taux depassee","Technique","Moyen","Resolu","450 EUR"),
    ("RS2","DPO : recommandation stockage images","Legal","Faible","Resolu","200 EUR"),
    ("RS3","Stripe : delai validation compte pro","Paiement","Eleve","En cours","1 200 EUR"),
    ("RS4","Derive charge Dev Mobile (S7)","RH","Moyen","En cours","2 800 EUR"),
]
for i, (rid,titre,cat,imp,stat,cout) in enumerate(risques):
    t.rows[i+1].cells[0].text = rid; t.rows[i+1].cells[1].text = titre
    t.rows[i+1].cells[2].text = cat; t.rows[i+1].cells[3].text = imp
    t.rows[i+1].cells[4].text = stat; t.rows[i+1].cells[5].text = cout

p("Cout total des risques survenus : 4 650 EUR (soit 43% de la provision de 10 800 EUR).")
p("Provision restante : 6 150 EUR — reserve preservee.")

# ============================================================
h1("7. PROJECTIONS FINANCIERES (ROI INVESTISSEURS)")
# ============================================================

h2("7.1 Investissement total")
p("Investissement = Budget RH + Hors-RH + Risques = ~130 000 EUR")

h2("7.2 Projections de CA mensuel (M1-M24)")
t = doc.add_table(rows=13, cols=3, style='Light Grid Accent 1')
t.rows[0].cells[0].text = "Mois"; t.rows[0].cells[1].text = "CA Prevu"; t.rows[0].cells[2].text = "CA Reel"
ca_data = [("M1-M2 (dev)","0 EUR","0 EUR"),("M3","1 500 EUR","1 200 EUR"),
    ("M4","3 000 EUR","2 800 EUR"),("M5","6 000 EUR","5 500 EUR"),
    ("M6","9 000 EUR","—"),("M7","12 000 EUR","—"),("M8","15 000 EUR","—"),
    ("M9","18 000 EUR","—"),("M10","20 000 EUR","—"),("M11","22 000 EUR","—"),
    ("M12","25 000 EUR","—"),("M13-M24","25 000 EUR/mois","—")]
for i, (m,cp,cr) in enumerate(ca_data):
    t.rows[i+1].cells[0].text = m; t.rows[i+1].cells[1].text = cp; t.rows[i+1].cells[2].text = cr

h2("7.3 Ratios investisseurs")
p("- Breakeven (seuil de rentabilite) : estime a M9-M10")
p("- OPEX mensuel post-lancement : 1 200 EUR/mois (maintenance, infra)")
p("- ROI a 24 mois : positif (CA cumule > Investissement + OPEX)")
p("- CA cumule prevu a M12 : ~137 500 EUR")
p("- CA cumule prevu a M24 : ~437 500 EUR")

# ============================================================
h1("8. PHASAGE BUDGETAIRE MENSUEL")
# ============================================================
p("Repartition des depenses sur les 2 mois de developpement :")
t = doc.add_table(rows=3, cols=3, style='Light Grid Accent 1')
t.rows[0].cells[0].text = ""; t.rows[0].cells[1].text = "Janvier 2025 (S1-S4)"
t.rows[0].cells[2].text = "Fevrier 2025 (S5-S8)"
t.rows[1].cells[0].text = "Poids RH"; t.rows[1].cells[1].text = "45%"; t.rows[1].cells[2].text = "55%"
t.rows[2].cells[0].text = "Logique"; t.rows[2].cells[1].text = "Cadrage + debut dev"
t.rows[2].cells[2].text = "Dev intensif + tests + production"

# ============================================================
h1("9. DASHBOARD STREAMLIT — PAGES DU SYSTEME")
# ============================================================
p("L'application de pilotage comporte 4 dashboards interactifs :")
p("1. Equipe Projet : Gantt interactif, charge par ressource, velocity, meteo equipe")
p("2. Budget Global : Repartition par categorie, phasage hebdomadaire, leviers d'optimisation, registre des risques")
p("3. Direction Financiere : KPIs jalons S4/S8, courbe S planifie vs reel, analyse par ressource, heatmap risques")
p("4. Investisseurs : Projections CA sur 24 mois, breakeven, ROI, phasage mensuel depenses vs revenus")

# ============================================================
h1("10. SYNTHESE CHIFFRES CLES")
# ============================================================
t = doc.add_table(rows=12, cols=2, style='Light Grid Accent 1')
t.rows[0].cells[0].text = "Indicateur"; t.rows[0].cells[1].text = "Valeur"
synth = [
    ("Budget total valide", "130 000 EUR"),
    ("Cout RH", "112 000 EUR (86.2%)"),
    ("Cout hors-RH", "7 200 EUR (5.5%)"),
    ("Provision risques", "10 800 EUR (8.3%)"),
    ("Cout reel a S8", "33 422 EUR"),
    ("Ecart budgetaire", "+5.8%"),
    ("Taches livrees", "24 / 24 (100%)"),
    ("Risques survenus", "4 / provision 10 800 EUR"),
    ("Cout risques", "4 650 EUR (43% provision)"),
    ("Breakeven", "M9-M10"),
    ("Equipe", "5 profils actifs / 9 initiaux"),
]
for i, (k,v) in enumerate(synth):
    t.rows[i+1].cells[0].text = k; t.rows[i+1].cells[1].text = v

# ============================================================
# SAUVEGARDE
# ============================================================
output = "BookOne_Donnees_NotebookLM.docx"
doc.save(output)
print(f"Document genere : {output}")
